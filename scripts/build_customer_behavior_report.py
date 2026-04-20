"""
거래처 거동(去動) 보고서 v1 — 4개 섹션 종합

1) 매출 붕괴 해부도 (축소·이탈·신규 한 장 요약)
2) 이탈 경보 룰 + 현재 축소 43개 위험 점수
3) 신규 거래처 정착 리뷰
4) 발행 패턴별 영업 관리 분리

입력: customer_monthly_matrix.json, customer_lifecycle.json, billing_pattern.json
출력: AX 전환 계획/YYYY-MM-DD/거래처_거동보고서_v1_YYYYMMDD.docx
      + AX 전환 계획/YYYY-MM-DD/거래처_거동보고서_v1_부록_YYYYMMDD.xlsx

원칙:
- 판정어("부실/🔴/소명 대상") 사용 금지 — 라포 쌓이기 전엔 질문형 언어
- 위험도는 "확인 필요 / 관찰 / 보고 중심"으로 3단계
"""
from __future__ import annotations

import json
import sys
from datetime import datetime
from pathlib import Path

try:
    import pandas as pd
    from docx import Document
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor, Cm
    from openpyxl import Workbook
    from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
    from openpyxl.utils import get_column_letter
except ImportError as e:
    print(f"❌ pip3 install python-docx pandas openpyxl: {e}")
    sys.exit(1)

SCRIPT_DIR = Path(__file__).parent
IN_MATRIX = SCRIPT_DIR / "output" / "customer_monthly_matrix.json"
IN_LIFECYCLE = SCRIPT_DIR / "output" / "customer_lifecycle.json"
IN_BILLING = SCRIPT_DIR / "output" / "billing_pattern.json"

TODAY = datetime.now().strftime("%Y-%m-%d")
OUT_DIR = Path(f"/Users/jack/dev/gabwoo/AX 전환 계획/{TODAY}")
OUT_DIR.mkdir(parents=True, exist_ok=True)
STAMP = datetime.now().strftime("%Y%m%d")
OUT_DOCX = OUT_DIR / f"거래처_거동보고서_v1_{STAMP}.docx"
OUT_XLSX = OUT_DIR / f"거래처_거동보고서_v1_부록_{STAMP}.xlsx"

MONTHS_36 = [f"{y}-{m:02d}" for y in (2023, 2024, 2025) for m in range(1, 13)]


# ───────── 데이터 로드 ─────────
def load_data():
    matrix = json.loads(IN_MATRIX.read_text())
    lc = json.loads(IN_LIFECYCLE.read_text())
    billing = json.loads(IN_BILLING.read_text())
    return matrix, lc["churn"], lc["new"], billing


# ───────── 1) 축소 43개 위험 점수 ─────────
def score_shrinkage(matrix):
    """축소(A) 거래처의 현재 위험도 계산.
    신호:
      R1: 2025 매출 / 2023 매출 비율 (낮을수록 위험)
      R2: 최근 3개월(2025-10,11,12) 공백 여부
      R3: 2024 하반기 vs 2025 하반기 감소율
      R4: 마지막 거래 월이 2025-06 이전
    """
    rows = []
    for c in matrix:
        if c["패턴"] != "A. 축소":
            continue
        monthly = c["월별"]
        s23, s25 = c["2023_합계"], c["2025_합계"]
        # R1
        r1 = (s25 / s23) if s23 > 0 else 1.0
        # R2: 마지막 3개월 공백
        last_3 = [monthly.get(m, 0) for m in ["2025-10", "2025-11", "2025-12"]]
        r2 = sum(1 for v in last_3 if v > 0)  # 0이면 3개월 공백
        # R3: 24H2 vs 25H2 비교
        h24 = sum(monthly.get(f"2024-{m:02d}", 0) for m in range(7, 13))
        h25 = sum(monthly.get(f"2025-{m:02d}", 0) for m in range(7, 13))
        r3 = (h25 / h24) if h24 > 0 else 1.0
        # R4: 마지막 거래 월
        positive = [(m, v) for m, v in monthly.items() if v > 0]
        last_m = positive[-1][0] if positive else None
        r4 = 1 if (last_m and last_m < "2025-07") else 0
        # 점수 (높을수록 위험)
        score = 0
        if r1 < 0.3: score += 3
        elif r1 < 0.5: score += 2
        elif r1 < 0.7: score += 1
        if r2 == 0: score += 3
        elif r2 == 1: score += 1
        if r3 < 0.3: score += 2
        elif r3 < 0.6: score += 1
        if r4 == 1: score += 2
        # 등급
        if score >= 6:
            grade = "확인 필요"
        elif score >= 3:
            grade = "관찰"
        else:
            grade = "보고 중심"
        rows.append({
            "회사": c["회사"],
            "거래처": c["거래처"],
            "2023(억)": round(s23 / 1e8, 2),
            "2024(억)": round(c["2024_합계"] / 1e8, 2),
            "2025(억)": round(s25 / 1e8, 2),
            "23→25비율": round(r1, 2),
            "24H2→25H2": round(r3, 2),
            "Q4_활동월수": r2,
            "마지막거래월": last_m or "-",
            "위험점수": score,
            "등급": grade,
        })
    df = pd.DataFrame(rows).sort_values(["위험점수", "2023(억)"], ascending=[False, False]).reset_index(drop=True)
    return df


# ───────── 2) 이탈 28개사에서 추출한 사후 신호 ─────────
def analyze_churn_signals(churn):
    """이탈 28개가 공통적으로 보인 "이탈 전" 신호를 재계산."""
    ratios = [c["마지막/전6평균"] for c in churn if c.get("마지막/전6평균") is not None]
    signal_summary = {
        "케이스수": len(churn),
        "급감형(마지막월이 전6평균 0.5배 미만)": sum(1 for r in ratios if r < 0.5),
        "완만감소형(0.5~1.0배)": sum(1 for r in ratios if 0.5 <= r < 1.0),
        "유지→단절형(1.0~1.5배)": sum(1 for r in ratios if 1.0 <= r < 1.5),
        "과대발행후단절형(1.5배 이상)": sum(1 for r in ratios if r >= 1.5),
    }
    # 이탈 시점 분포
    timing = {}
    for c in churn:
        timing[c["이탈시점"]] = timing.get(c["이탈시점"], 0) + 1
    return signal_summary, timing


# ───────── Word 작성 헬퍼 ─────────
def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def add_h1(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(20)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)


def add_h2(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(15)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)
    # 섹션 구분용 상단 여백
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)


def add_h3(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(3)


def add_p(doc, text, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for r in p.runs:
        r.font.size = Pt(10)


def add_table(doc, headers, rows, widths=None, emphasize_first=True):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = "Light Grid Accent 1"
    # 헤더
    for j, h in enumerate(headers):
        cell = tbl.rows[0].cells[j]
        cell.text = str(h)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1F3864")
    # 본문
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = tbl.rows[i].cells[j]
            if isinstance(v, float):
                cell.text = f"{v:,.2f}" if abs(v) < 100 else f"{v:,.1f}"
            elif isinstance(v, int):
                cell.text = f"{v:,}"
            else:
                cell.text = str(v) if v is not None else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                if j > 0 and isinstance(v, (int, float)):
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if widths:
        for j, w in enumerate(widths):
            for row in tbl.rows:
                row.cells[j].width = Cm(w)


# ───────── 보고서 본문 ─────────
def build_docx(matrix, churn, new, billing, shrink_risk, signals, timing):
    doc = Document()
    # 기본 폰트
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(11)

    # 표지
    add_h1(doc, "갑우그룹 거래처 거동(去動) 보고서 v1")
    add_p(doc, f"작성일: {TODAY}  /  출처: 제품매출현황 xls 2개 파일 (2023·2024·2025 총 11,176행)", italic=True, size=10, color=(0x55, 0x55, 0x55))
    add_p(doc, "본 보고서는 세금계산서 발행 기준 실매출을 월 단위로 36개월 재구성해 " \
              "거래처별 '언제 빠지기 시작했는지', '누가 새로 들어왔는지', '어떻게 청구하는지'를 정리한 문서입니다. " \
              "수치는 질문을 생성하기 위한 것이며, 판정이 아닙니다.", size=10, color=(0x33, 0x33, 0x33))

    # ── 섹션 1: 매출 붕괴 해부도 ──
    add_h2(doc, "1. 매출 붕괴 해부도 — 2023 481억 → 2025 328억 (-153억, -31.7%)")
    add_p(doc, "1억+ 거래처 122개를 5개 패턴으로 분류한 결과, 매출 감소는 '축소 43개사'에 집중됩니다. " \
              "이탈·축소 손실 -257억을 신규·성장 +113억이 부분 상쇄하지만, 순 -145억의 상당 부분은 기존 거래처의 수량·단가 축소에서 발생했습니다.", size=10)

    add_h3(doc, "패턴 요약 (1억 이상 거래처)")
    add_table(doc,
        ["패턴", "개수", "2023(억)", "2024(억)", "2025(억)", "23→25 Δ(억)"],
        [
            ["A. 축소", 43, 272.0, None, 89.4, -182.6],
            ["C. 이탈", 28, 74.8, None, 0.0, -74.8],
            ["G. 성장", 19, 35.4, None, 65.4, +30.0],
            ["N. 신규", 20, 0.0, None, 82.9, +82.9],
            ["S. 유지", 12, 98.9, None, 90.6, -8.3],
            ["합계", 122, 481.1, None, 328.3, -152.8],
        ],
    )
    add_p(doc, "2024년 수치는 섹션 2~3 상세 표에 포함되어 있습니다.", italic=True, size=9, color=(0x55, 0x55, 0x55))

    add_h3(doc, "축소 43개사 상위 10 — 어디서 가장 많이 빠졌는가")
    shrink_sorted = sorted(
        [c for c in matrix if c["패턴"] == "A. 축소"],
        key=lambda x: x["2023_합계"] - x["2025_합계"], reverse=True,
    )[:10]
    rows = []
    for c in shrink_sorted:
        rows.append([
            c["회사"], c["거래처"][:18],
            round(c["2023_합계"] / 1e8, 1),
            round(c["2024_합계"] / 1e8, 1),
            round(c["2025_합계"] / 1e8, 1),
            round((c["2023_합계"] - c["2025_합계"]) / 1e8, 1),
        ])
    add_table(doc, ["회사", "거래처", "2023", "2024", "2025", "23→25 손실"], rows)

    add_p(doc, "※ 상위 10개사에서만 -124억이 빠졌습니다 (전체 축소 손실의 68%). " \
              "이 10곳에 대한 '수량인가 단가인가 이탈인가' 분리 진단이 가장 파급력이 큽니다.", size=10)

    # ── 섹션 2: 이탈 경보 룰 ──
    add_h2(doc, "2. 이탈 경보 룰 — 2024년 다 떠났다 + 현재 축소 43개 위험군")
    add_p(doc, "이탈 28개사 전부가 2024년 안에 떠났습니다. 2025년 신규 이탈은 0건입니다. " \
              "즉 2025년은 '이탈의 결과'이지 '이탈 진행 중'이 아닙니다. 남은 과제는 축소 43개사가 2026년 이탈로 이어지지 않게 막는 것입니다.", size=10)

    add_h3(doc, "이탈 28개사 시점 분포")
    add_table(doc,
        ["이탈 시점", "거래처 수"],
        [[k, v] for k, v in timing.items()] + [["2025년 이탈", 0]],
    )

    add_h3(doc, "이탈 직전 신호 분석 — 조기경보는 '감소율'보다 '공백 개월'이 더 강함")
    add_table(doc,
        ["신호 유형", "케이스 수", "해석"],
        [
            ["급감형 (마지막월 < 전6평균 0.5배)", signals["급감형(마지막월이 전6평균 0.5배 미만)"], "명확한 감소 경고 — 소수(6/28)"],
            ["완만감소형 (0.5~1.0배)", signals["완만감소형(0.5~1.0배)"], "서서히 빠진 케이스"],
            ["유지→단절형 (1.0~1.5배)", signals["유지→단절형(1.0~1.5배)"], "끝까지 정상 발행 후 갑자기 단절"],
            ["과대발행후단절형 (1.5배 이상)", signals["과대발행후단절형(1.5배 이상)"], "마지막 달 몰아 발행 후 이탈 — 주의 필요"],
        ],
    )
    add_p(doc, "→ 전체 28개 중 12개(43%)는 마지막 월까지 '정상 또는 과대'로 발행하다 갑자기 끊겼습니다. " \
              "즉 감소율 하나만으로는 이탈 예측이 어렵습니다. 더 강한 신호는 ①연간 매출 대폭 감소 + ②공백 개월 + ③분기 단절 조합입니다.", size=10)

    add_h3(doc, "현재 축소 43개사 위험 점수 (룰 기반)")
    add_p(doc, "룰: 23→25비율 / 24H2→25H2 감소율 / 2025 Q4 활동 월수 / 마지막 거래월을 가중 합산(0~10점).", size=10)
    risk_top = shrink_risk.head(15).copy()
    rows = []
    for _, r in risk_top.iterrows():
        rows.append([
            r["회사"], r["거래처"][:16],
            r["2023(억)"], r["2025(억)"],
            r["23→25비율"], r["24H2→25H2"],
            r["Q4_활동월수"], r["마지막거래월"],
            r["위험점수"], r["등급"],
        ])
    add_table(doc,
        ["회사", "거래처", "2023", "2025", "23→25", "24H2→25H2", "Q4활동", "마지막거래", "점수", "등급"],
        rows,
    )

    risk_counts = shrink_risk["등급"].value_counts().to_dict()
    add_p(doc, f"축소 43개 중 · 확인 필요 {risk_counts.get('확인 필요', 0)}개 · 관찰 {risk_counts.get('관찰', 0)}개 · 보고 중심 {risk_counts.get('보고 중심', 0)}개. " \
              "'확인 필요' 구간은 다음 분기 전에 담당 영업자에게 현장 상황을 물어볼 대상입니다.", size=10)

    # ── 섹션 3: 신규 정착 리뷰 ──
    add_h2(doc, "3. 신규 거래처 정착 리뷰 — 20개 중 13개는 정착, 7개는 아직")
    add_p(doc, "2023년에 없던 거래처 중 2025년에 1억 이상을 기록한 20개사입니다. 정착 점수는 최근 6개월 활동월수·매출비중을 50:50으로 합산했습니다.", size=10)

    # 정착 등급 요약
    grade_count = {"🟢 정착": 0, "🟡 관찰": 0, "🟠 확인": 0}
    for n in new:
        grade_count[n["정착등급"]] = grade_count.get(n["정착등급"], 0) + 1
    add_table(doc,
        ["정착 등급", "거래처 수", "2025 합(억)"],
        [
            ["정착 (🟢)", grade_count.get("🟢 정착", 0), round(sum(x["2025(억)"] for x in new if x["정착등급"]=="🟢 정착"), 1)],
            ["관찰 (🟡)", grade_count.get("🟡 관찰", 0), round(sum(x["2025(억)"] for x in new if x["정착등급"]=="🟡 관찰"), 1)],
            ["확인 (🟠)", grade_count.get("🟠 확인", 0), round(sum(x["2025(억)"] for x in new if x["정착등급"]=="🟠 확인"), 1)],
        ],
    )

    add_h3(doc, "신규 20개사 전체 — 진입시점·최근활동·정착점수")
    new_sorted = sorted(new, key=lambda x: x["3년(억)"], reverse=True)
    rows = []
    for n in new_sorted:
        rows.append([
            n["회사"], n["거래처"][:18], n["진입시점"],
            n["2024(억)"], n["2025(억)"],
            n["최근6개월_활동월수"], n["최근6개월매출(억)"],
            n["정착점수"], n["정착등급"],
        ])
    add_table(doc,
        ["회사", "거래처", "진입", "2024", "2025", "최근6M 활동", "최근6M 매출", "점수", "등급"],
        rows,
    )
    add_p(doc, "→ 신규 +82.9억은 이탈 -74.8억을 거의 상쇄합니다. 다만 '축소 43개 -182억' 손실은 덮지 못합니다. " \
              "신규 거래처가 2년차에 유지·성장으로 넘어가는 비율이 2026년 매출 회복의 핵심 변수입니다.", size=10)

    # ── 섹션 4: 발행 패턴 가이드 ──
    add_h2(doc, "4. 발행 패턴별 영업·재경 운영 가이드")
    add_p(doc, "1억+ 거래처 161개를 세금계산서 발행 방식으로 분류했습니다. 같은 매출이어도 '월합산'과 '건별 집중'은 캐시플로우 리스크와 영업 관리 룰이 다릅니다.", size=10)

    pattern_count = {}
    pattern_sum = {}
    for b in billing:
        p = b["발행패턴"]
        pattern_count[p] = pattern_count.get(p, 0) + 1
        pattern_sum[p] = pattern_sum.get(p, 0) + b["3년_합계(억)"]
    add_table(doc,
        ["발행 패턴", "거래처 수", "3년 합(억)", "운영 함의"],
        [
            ["월합산", pattern_count.get("월합산", 0), round(pattern_sum.get("월합산", 0), 1),
             "연 24건 이하 + 활동 8개월+. 월마감 캐시플로우 리스크 큰 군. 회수일 관리 필요"],
            ["건별 혼합", pattern_count.get("건별 혼합", 0), round(pattern_sum.get("건별 혼합", 0), 1),
             "일반적 건별 청구. 수주 파이프라인 가시성 높음"],
            ["건별 집중", pattern_count.get("건별 집중", 0), round(pattern_sum.get("건별 집중", 0), 1),
             "연 100건+ 소량 빈번 거래. 단가·인당 마진 관리 대상"],
            ["이월·간헐", pattern_count.get("이월·간헐", 0), round(pattern_sum.get("이월·간헐", 0), 1),
             "활동 5개월 이하. 교원 사급·프로젝트성 거래 후보. 예측모델에서 제외해야 왜곡 없음"],
        ],
    )

    add_h3(doc, "이월·간헐 53개사 중 3년합 상위 10 — 사급 구조 후보")
    ig = [b for b in billing if b["발행패턴"] == "이월·간헐"]
    ig_sorted = sorted(ig, key=lambda x: x["3년_합계(억)"], reverse=True)[:10]
    rows = []
    for b in ig_sorted:
        rows.append([
            b["회사"], b["거래처"][:18], b["패턴"],
            b["평균연발행건수"], b["평균활동월수"],
            b["건당평균(백만)"], b["3년_합계(억)"],
        ])
    add_table(doc,
        ["회사", "거래처", "패턴", "연발행건수", "활동월수", "건당(백만)", "3년(억)"],
        rows,
    )
    add_p(doc, "→ 교원·이투스 등 지입거래처와 연말·연초 일괄 처리 패턴이 섞여 있습니다. " \
              "이 53개사는 월별 대시보드 예측치에서 별도 트랙으로 관리하면 월 편차 해석이 훨씬 깨끗해집니다.", size=10)

    # 마무리
    add_h2(doc, "5. 다음 액션 — 우선순위 4가지")
    add_bullet(doc, "① (이번 주) 축소 43개 중 '확인 필요' 등급 거래처에 담당 영업자 현장 체크 요청")
    add_bullet(doc, "② (이번 달) 신규 20개사 중 🟡 관찰·🟠 확인 7개사 담당자와 2026 계획 점검")
    add_bullet(doc, "③ (다음 달) 이월·간헐 53개사 중 상위 10개 사급 구조 확인 → 대시보드 분리 뷰 구현")
    add_bullet(doc, "④ (상시) 이탈 경보 룰을 월간 배치로 돌려 축소 거래처 상태 자동 감시")

    doc.save(OUT_DOCX)
    print(f"✅ docx 저장: {OUT_DOCX}")


# ───────── 부록 xlsx ─────────
HEADER_FILL = PatternFill("solid", fgColor="1F3864")
HEADER_FONT = Font(bold=True, color="FFFFFF", size=11)
BORDER = Border(
    left=Side(style="thin", color="BFBFBF"),
    right=Side(style="thin", color="BFBFBF"),
    top=Side(style="thin", color="BFBFBF"),
    bottom=Side(style="thin", color="BFBFBF"),
)


def write_xlsx(shrink_risk, new, billing, matrix):
    wb = Workbook()
    wb.active.title = "1_축소43_위험점수"
    ws1 = wb["1_축소43_위험점수"]
    _write(ws1, shrink_risk, "축소 43개사 위험 점수 — 확인 필요 / 관찰 / 보고 중심")
    ws2 = wb.create_sheet("2_신규20_정착")
    new_df = pd.DataFrame(new).sort_values("3년(억)", ascending=False)
    _write(ws2, new_df, "신규 20개사 정착도")
    ws3 = wb.create_sheet("3_발행패턴_161개")
    bill_df = pd.DataFrame(billing)
    _write(ws3, bill_df, "1억+ 거래처 발행 패턴 분류")
    ws4 = wb.create_sheet("4_이월간헐_53개")
    ig_df = bill_df[bill_df["발행패턴"] == "이월·간헐"].copy()
    _write(ws4, ig_df, "이월·간헐 53개사 (사급 구조 후보)")
    wb.save(OUT_XLSX)
    print(f"✅ xlsx 부록 저장: {OUT_XLSX}")


def _write(ws, df, title):
    ws["A1"] = title
    ws["A1"].font = Font(bold=True, color="FFFFFF", size=12)
    ws["A1"].fill = PatternFill("solid", fgColor="2E75B6")
    ncols = max(len(df.columns), 5)
    ws.merge_cells(start_row=1, start_column=1, end_row=1, end_column=ncols)
    r = 3
    for j, col in enumerate(df.columns, start=1):
        c = ws.cell(row=r, column=j, value=str(col))
        c.font = HEADER_FONT
        c.fill = HEADER_FILL
        c.border = BORDER
        c.alignment = Alignment(horizontal="center")
    for i, row in enumerate(df.itertuples(index=False), start=r + 1):
        for j, val in enumerate(row, start=1):
            c = ws.cell(row=i, column=j, value=val)
            c.border = BORDER
            if isinstance(val, (int, float)) and not isinstance(val, bool):
                c.alignment = Alignment(horizontal="right")
                if isinstance(val, float):
                    c.number_format = "#,##0.00"
                else:
                    c.number_format = "#,##0"
    for col in ws.columns:
        max_len = 0
        col_letter = None
        for cell in col:
            if col_letter is None:
                col_letter = get_column_letter(cell.column)
            v = "" if cell.value is None else str(cell.value)
            ln = sum(2 if ord(c) > 127 else 1 for c in v)
            if ln > max_len:
                max_len = ln
        if col_letter:
            ws.column_dimensions[col_letter].width = min(max_len + 2, 50)


def main():
    matrix, churn, new, billing = load_data()
    print(f"📊 matrix={len(matrix)} churn={len(churn)} new={len(new)} billing={len(billing)}")
    shrink_risk = score_shrinkage(matrix)
    print(f"축소 43개 등급: {shrink_risk['등급'].value_counts().to_dict()}")
    signals, timing = analyze_churn_signals(churn)
    print(f"이탈 신호: {signals}")
    print(f"이탈 시점: {timing}")
    build_docx(matrix, churn, new, billing, shrink_risk, signals, timing)
    write_xlsx(shrink_risk, new, billing, matrix)


if __name__ == "__main__":
    main()
